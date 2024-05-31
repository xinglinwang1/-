import os
import json
from docx import Document
import re

def extract_text_from_paragraph(paragraph):
    return paragraph.text.strip()

def get_next_non_empty_paragraph(doc, current_index):
    next_index = current_index + 1
    while next_index < len(doc.paragraphs):
        text = extract_text_from_paragraph(doc.paragraphs[next_index])
        if text:
            return text
        next_index += 1
    return ""

def process_docx(file_path):
    doc = Document(file_path)
    # 定义一个字典，以键值对的形式存储各级标题和内容
    # {
    #    一级标题：{
    #       二级标题：{
    #          三级标题：内容
    #       }
    #    }
    # }
    content = {}
    current_section = ""
    current_subsection = ""
    current_subsubsection = ""

    for index, paragraph in enumerate(doc.paragraphs):
        text = extract_text_from_paragraph(paragraph)
        if not text:
            continue

        # Detect section headers (e.g., "1 总则" or "1总则")
        if re.match(r"^\d+(?!\.)\s*(.*)", text):
            current_section = re.match(r"^\d+(?!\.)\s*(.*)", text).group(1)
            next_text = get_next_non_empty_paragraph(doc, index)
            if re.match(r"^\d+\.\d+(?!\.)\s*(.*)", next_text):
                content[current_section] = {}
            else:
                content[current_section] = ""
            current_subsection = ""
            current_subsubsection = ""
        # Detect subsection headers (e.g., "1.1 编制目的" or "1.1编制目的")
        elif re.match(r"^\d+\.\d+(?!\.)\s*(.*)", text):
            current_subsection = re.match(r"^\d+\.\d+(?!\.)\s*(.*)", text).group(1)
            if current_section:
                if isinstance(content[current_section], str):
                    content[current_section] = {}
                next_text = get_next_non_empty_paragraph(doc, index)
                if re.match(r"^\d\.\d\.\d+\s*(.*)", next_text):
                    content[current_section][current_subsection] = {}
                else:
                    content[current_section][current_subsection] = ""
            current_subsubsection = ""
        # Detect subsubsection headers (e.g., "1.1.1 气象水文海洋信息" or "1.1.1气象水文海洋信息")
        elif re.match(r"^\d\.\d\.\d+\s*(.*)", text):
            current_subsubsection = re.match(r"^\d\.\d\.\d+\s*(.*)", text).group(1)
            if current_section and current_subsection:
                if isinstance(content[current_section][current_subsection], str):
                    content[current_section][current_subsection] = {}
                content[current_section][current_subsection][current_subsubsection] = ""


        # # Detect section headers (e.g., "1 总则")
        # if re.match(r'^\d+ ', text):
        #     current_section = text.split(' ', 1)[1]
        #     next_text = get_next_non_empty_paragraph(doc, index)
        #     if re.match(r'^\d+\.\d+ ', next_text):
        #         content[current_section] = {}
        #     else:
        #         content[current_section] = ""
        #     current_subsection = ""
        #     current_subsubsection = ""
        #
        # # Detect subsection headers (e.g., "1.1 编制目的")
        # elif re.match(r'^\d+\.\d+ ', text):
        #     current_subsection = text.split(' ', 1)[1]
        #     if current_section:
        #         next_text = get_next_non_empty_paragraph(doc, index)
        #         if re.match(r'^\d+\.\d+\.\d+ ', next_text):
        #             content[current_section][current_subsection] = {}
        #         else:
        #             content[current_section][current_subsection] = ""
        #     current_subsubsection = ""
        #
        # # Detect subsubsection headers (e.g., "1.1.1 气象水文海洋信息")
        # elif re.match(r'^\d+\.\d+\.\d+ ', text):
        #     current_subsubsection = text.split(' ', 1)[1]
        #     if current_section and current_subsection:
        #         content[current_section][current_subsection][current_subsubsection] = ""



        # Detect content
        else:
            if current_section and current_subsection and current_subsubsection:
                content[current_section][current_subsection][current_subsubsection] += text
            elif current_section and current_subsection:
                if current_subsection in content[current_section]:
                    content[current_section][current_subsection] += text
                else:
                    content[current_section][current_subsection] = text
            elif current_section:
                if current_section in content:
                    content[current_section] += text
                else:
                    content[current_section] = text
        #print(content)

    return content


def format_to_key_value_pairs(content, file_name):
    key_value_pairs = []
    for section, subsections in content.items():
        # Check if subsections is a string (which means there are no further nested sections)
        if isinstance(subsections, str):
            input_text = f"请编制{file_name}当中的{section}"
            output_text = subsections.strip()
            key_value_pairs.append({"input": input_text, "output": output_text})
        else:
            for subsection, subsubsections in subsections.items():
                # Check if subsubsections is a string (which means there are no further nested sections)
                if isinstance(subsubsections, str):
                    input_text = f"请编制{file_name}当中的{section}当中的{subsection}"
                    output_text = subsubsections.strip()
                    key_value_pairs.append({"input": input_text, "output": output_text})
                else:
                    for subsubsection, text in subsubsections.items():
                        input_text = f"请编制{file_name}当中的{section}当中的{subsection}当中的{subsubsection}"
                        output_text = text.strip()
                        key_value_pairs.append({"input": input_text, "output": output_text})
    return key_value_pairs


def save_to_json(key_value_pairs, output_file):
    with open(output_file, 'w', encoding='utf-8') as f:
        for pair in key_value_pairs:
            json.dump(pair, f, ensure_ascii=False)
            f.write("\n")


def process_directory(directory_path):
    for file_name in os.listdir(directory_path):
        if file_name.endswith('.docx'):
            file_name1 = os.path.splitext(os.path.basename(file_name))[0]
            file_path = os.path.join(directory_path, file_name)
            content = process_docx(file_path)
            key_value_pairs = format_to_key_value_pairs(content, file_name1)
            output_file = os.path.join(directory_path, file_name.replace('.docx', '.json'))
            save_to_json(key_value_pairs, output_file)


if __name__ == "__main__":
    directory_path = r"C:\Users\wangxinglin\Desktop\test"
    process_directory(directory_path)